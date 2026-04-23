import os
import re
import sys
import uuid
from datetime import datetime

from docx import Document
from dotenv import load_dotenv
from pymongo import MongoClient

load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))

TP_TITLES = {
    "Secrétaire Assistant": "SA",
    "Assistant de Direction": "AD",
    "Conseiller Commercial": "CC",
    "Négociateur Technico-Commercial": "NTC",
    "Responsable d'Établissement Marchand": "REM",
}

SCHOOL_LEVELS = {
    "CAP / BEP avec 1 an": "CAP_BEP_WITH_1Y_EXP",
    "première / terminale avec 1 an": "PREMIERE_TERMINALE_WITH_1Y_EXP",
    "première / terminale": "PREMIERE_TERMINALE",
    "BAC avec 1 an": "BAC_WITH_1Y_EXP",
    "BAC+2 et plus": "BAC_PLUS_2_PLUS",
    "BAC+2": "BAC_PLUS_2",
    "BAC+3": "BAC_PLUS_3_PLUS",
    "BAC et plus": "BAC_PLUS",
    "BAC": "BAC",
}

DISCOVERY_SOURCES = {
    "Réseaux sociaux": "SOCIAL_MEDIA",
    "France Travail": "FRANCE_TRAVAIL",
    "Mission Locale": "MISSION_LOCALE",
    "Bouche à oreille": "WORD_OF_MOUTH",
    "Koann": "KOANN",
    "Autres": "OTHER",
}

TRAINING_SITES = {
    "Nord": "NORD_SAINTE_MARIE",
    "Ouest": "OUEST_SAINT_PAUL",
    "Sud": "SUD_SAINT_PIERRE",
}


def get_mongo_collection():
    username = os.getenv("MONGO_ROOT_USERNAME")
    password = os.getenv("MONGO_ROOT_PASSWORD")
    port = os.getenv("MONGO_PORT", "27017")
    client = MongoClient(
        f"mongodb://{username}:{password}@localhost:{port}/human_ressources?authSource=admin"
    )
    return client["human_ressources"]["candidates"]


def detect_tp_type(doc):
    for p in doc.paragraphs[:5]:
        for title, code in TP_TITLES.items():
            if title in p.text:
                return code
    return None


def inline_value(paragraph):
    """Return (label, value) where label=bold runs, value=non-bold runs."""
    label_parts, value_parts = [], []
    for run in paragraph.runs:
        (label_parts if run.bold else value_parts).append(run.text)
    label = "".join(label_parts).strip("\n\xa0 :")
    value = "".join(value_parts).strip("\n\xa0 ")
    if not value and ":" in label:
        label, _, value = label.partition(":")
        value = value.strip()
    return label.strip(), value.strip()


def is_selected(paragraph):
    """True if any run in the paragraph is underlined or bold-and-not-label."""
    for run in paragraph.runs:
        if run.underline:
            return True
    return False


def parse_yes_no(paragraphs, start, window=3):
    """Look ahead up to `window` paragraphs for Oui/Non selection."""
    oui, non = None, None
    for p in paragraphs[start : start + window]:
        t = p.text.strip()
        if t.startswith("Oui"):
            oui = p
        elif t.startswith("Non"):
            non = p
    if oui and is_selected(oui):
        return True
    if non and is_selected(non):
        return False
    return None


def extract_numbered_list(text, count=3):
    """Extract up to `count` numbered items (1. / 2. / 3.) from text."""
    items = re.findall(r"\d\.\s*(.+?)(?=\s*\d\.|$)", text, re.DOTALL)
    return [i.strip() for i in items[:count] if i.strip()]


def extract_skills(doc):
    for table in doc.tables:
        if not table.rows:
            continue
        header = table.rows[0].cells[0].text.strip()
        if "Compétences professionnelles" in header and len(table.columns) == 2:
            skills = []
            for row in table.rows[1:]:
                competence = row.cells[0].text.strip()
                raw_level = row.cells[1].text.strip()
                if competence:
                    level = raw_level if raw_level in ("A", "ECA", "NA", "NE") else "NE"
                    skills.append({"competence": competence, "level": level})
            return skills
    return []


def extract_sectors_and_company_skills(doc):
    sectors, company_skills = [], []
    for table in doc.tables:
        if not table.rows:
            continue
        header = table.rows[0].cells[0].text.strip()
        if "Secteurs d'activité" in header and len(table.columns) == 2:
            for row in table.rows[1:]:
                left = row.cells[0].text.strip()
                right = row.cells[1].text.strip()
                if left:
                    for p in row.cells[0].paragraphs:
                        if is_selected(p):
                            sectors.append(left)
                            break
                if right:
                    for p in row.cells[1].paragraphs:
                        if is_selected(p):
                            company_skills.append(right)
                            break
            break
    return sectors, company_skills


def parse_docx(file_path):
    doc = Document(file_path)
    paragraphs = doc.paragraphs
    tp_type = detect_tp_type(doc)

    identity = {}
    education = {}
    training_site = None
    support = {}
    immersion_agreement = None
    background = {"professional_experiences": []}
    profile = {}
    professional_projects = {}
    job_info = {}
    synthesis = {"pedagogical_recommendations": {}}

    in_experiences = False
    exp_buffer = []
    previous_formation = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            in_experiences = False
            if exp_buffer:
                background["professional_experiences"].append(
                    {"responsibilities": " ".join(exp_buffer)}
                )
                exp_buffer = []
            continue

        label, value = inline_value(p)
        lower: str = text.lower()
        if previous_formation == True:
            background["previous_trainings"] = lower
            previous_formation = False
            continue


        # Identity
        if "nom(s) et prénom(s) complets" in lower:
            if value:
                identity["full_name"] = value
        elif "date et lieu de naissance" in lower:
            if value:
                # Keep raw; structured parsing would need locale-aware date parsing
                identity["place_of_birth"] = value
        elif re.search(r"^âge\s*:", lower):
            if value:
                try:
                    identity["age"] = int(re.sub(r"\D", "", value))
                except ValueError:
                    pass
        elif "code" in lower and "postal" in lower:
            if value:
                parts = value.split(" ", 1)
                identity["postal_code"] = parts[0]
                if len(parts) > 1:
                    identity["city"] = parts[1].split(':')[1]
        # elif re.search(r"^email\s*:", lower):
        elif "email" in lower:
            identity["email"] = lower.split(':')[1].replace(' ', '')
        elif "numéro de téléphone" in lower:
            if value:
                identity["phone"] = re.sub('\D', '', value)
        elif "permis b" in lower:
            identity["driving_license_b"] = parse_yes_no(paragraphs, i + 1) or (
                value.lower().startswith("oui") if value else None
            )
        elif "moyen de transport" in lower:
            if value:
                identity["transport_means"] = value.split(":")[1]
        elif "référent psh" in lower:
            identity["psh_referral_request"] = parse_yes_no(paragraphs, i + 1)

        # Education — school level (option paragraphs are underlined when selected)
        elif "cap / bep" in lower and "expérience" in lower:
            if is_selected(p):
                education["school_level"] = "CAP_BEP_WITH_1Y_EXP"
        elif "première / terminale" in lower and "1 an" in lower:
            if is_selected(p):
                education["school_level"] = "PREMIERE_TERMINALE_WITH_1Y_EXP"
        elif "première / terminale" in lower:
            if is_selected(p):
                education["school_level"] = "PREMIERE_TERMINALE"
        elif "bac avec 1 an" in lower:
            if is_selected(p):
                education["school_level"] = "BAC_WITH_1Y_EXP"
        elif "bac+3" in lower or "bac+2 et plus" in lower:
            if is_selected(p):
                education["school_level"] = "BAC_PLUS_2_PLUS"
        elif "bac+2" in lower:
            if is_selected(p):
                education["school_level"] = "BAC_PLUS_2"
        elif re.search(r"bac et plus", lower):
            if is_selected(p):
                education["school_level"] = "BAC_PLUS"
        elif re.search(r"^niveau bac\b", lower):
            if is_selected(p):
                education["school_level"] = "BAC"
        elif "justificatif" in lower:
            if value:
                education["justification"] = value

        # Training site
        elif "disciplina nord" in lower:
            if is_selected(p):
                training_site = "NORD_SAINTE_MARIE"
        elif "disciplina ouest" in lower:
            if is_selected(p):
                training_site = "OUEST_SAINT_PAUL"
        elif "disciplina sud" in lower:
            if is_selected(p):
                training_site = "SUD_SAINT_PIERRE"

        # Support
        elif "france travail" in lower and "inscrit" in lower:
            registered = parse_yes_no(paragraphs, i + 1)
            support["france_travail_registered"] = registered
            # Agency is on the same line after "→ Si oui, dans quelle agence :"
            agency_match = re.search(r"agence\s*:\s*(.+)", text, re.IGNORECASE)
            if agency_match:
                support["france_travail_agency"] = agency_match.group(1).strip()
        elif "mission locale" in lower and "inscrit" in lower:
            registered = parse_yes_no(paragraphs, i + 1)
            support["mission_locale_registered"] = registered
            city_match = re.search(r"ville\s*:\s*(.+)", text, re.IGNORECASE)
            if city_match:
                support["mission_locale_city"] = city_match.group(1).strip()

        # Immersion
        elif "immersion en entreprise" in lower:
            immersion_agreement = parse_yes_no(paragraphs, i + 1)

        # Background
        elif "dernier diplôme" in lower:
            if value:
                background["last_diploma"] = value.split(":")[1]
        elif "formations suivies" in lower:
            previous_formation = True
        elif "expériences professionnelles passées" in lower:
            in_experiences = True

        # Profile
        elif "français (note sur 10)" in lower:
            if value:
                try:
                    profile["french_level"] = int(re.sub(r"\D", "", value[value.find(")"):]))
                except ValueError:
                    pass
        elif "anglais (note sur 10)" in lower:
            if value:
                try:
                    profile["english_level"] = int(re.sub(r"\D", "", value))
                except ValueError:
                    pass
        elif "autres langues" in lower:
            if value:
                profile["other_languages"] = [l.strip() for l in value.split(",") if l.strip()]
        elif re.search(r"^qualités\s*:", lower):
            items = extract_numbered_list(text)
            if items:
                profile["qualities"] = items
        elif re.search(r"^défauts\s*:", lower):
            items = extract_numbered_list(text)
            if items:
                profile["defects"] = items
        elif "compétences numériques" in lower:
            if value:
                profile["digital_skills"] = [s.strip() for s in value.split(",") if s.strip()]
        elif "relever des défis" in lower:
            profile["ready_for_challenges"] = parse_yes_no(paragraphs, i + 1)
        elif "passions ou hobbies" in lower:
            if value:
                profile["hobbies"] = value
        elif "points forts" in lower and "axes d'amélioration" in lower:
            if value:
                profile["strengths_and_improvements"] = value

        # Professional projects
        elif "objectifs professionnels" in lower:
            if value:
                professional_projects["career_objectives"] = value.split(':')[1]
        elif "compétences que vous souhaitez" in lower:
            if value:
                professional_projects["desired_skills"] = value.split(':')[1]
        elif "motivations pour choisir" in lower:
            if value:
                professional_projects["apprenticeship_motivation"] = value.split(':')[1]
        elif "attentes vis-à-vis" in lower:
            print(f"value: [{value}], label: [{label}], lower: [{lower}]")
            if value:
                professional_projects["training_expectations"] = lower.split(':')[1]

        # Job info
        elif "pourquoi avez-vous choisi ce domaine" in lower:
            if value:
                job_info["domain_motivation"] = value.split('?')[1]
        elif "questions / préoccupations" in lower:
            if value:
                job_info["questions_concerns"] = value.split('?')[1]
        elif "date de disponibilité" in lower:
            if value:
                job_info["availability_date_raw"] = value.split('?')[1]
        elif "mobilité géographique" in lower:
            if value:
                job_info["geographic_mobility"] = value.split('?')[1]
        elif "week-end" in lower:
            job_info["weekend_work"] = parse_yes_no(paragraphs, i + 1) or (
                value.lower().startswith("oui") if value else None
            )

        # Discovery source
        elif text in DISCOVERY_SOURCES and is_selected(p):
            job_info["discovery_source"] = DISCOVERY_SOURCES[text]

        # Synthesis
        elif "conclusion sur la faisabilité" in lower:
            if value:
                synthesis["feasibility_conclusion"] = value.split(':')[1]
        elif "pertinence du parcours" in lower:
            if value:
                synthesis["pathway_relevance"] = value.split(':')[1]
        elif "besoins particuliers" in lower:
            if value:
                synthesis["special_needs"] = value
        elif "renforcement en bureautique" in lower:
            synthesis["pedagogical_recommendations"]["office_tools_reinforcement"] = is_selected(p)
        elif "communication écrite" in lower:
            synthesis["pedagogical_recommendations"]["written_communication_support"] = is_selected(p)
        elif "confiance à l'oral" in lower:
            synthesis["pedagogical_recommendations"]["oral_confidence_development"] = is_selected(p)
        elif "gestion du temps" in lower and "organisation" in lower:
            synthesis["pedagogical_recommendations"]["time_management_support"] = is_selected(p)
        elif "posture professionnelle" in lower:
            synthesis["pedagogical_recommendations"]["professional_posture_work"] = is_selected(p)
        elif "formation complémentaire en langue" in lower:
            synthesis["pedagogical_recommendations"]["language_training"] = is_selected(p)
        elif "immersion renforcée" in lower:
            synthesis["pedagogical_recommendations"]["enhanced_company_immersion"] = is_selected(p)
        elif "accompagnement spécifique psh" in lower:
            synthesis["pedagogical_recommendations"]["psh_specific_support"] = is_selected(p)
        elif "suivi individualisé" in lower:
            synthesis["pedagogical_recommendations"]["individual_follow_up"] = is_selected(p)
        elif "gestion du stress" in lower and "suivi" in lower:
            synthesis["pedagogical_recommendations"]["stress_management_follow_up"] = is_selected(p)
        elif "autres préconisations" in lower:
            if value:
                synthesis["other_recommendations"] = value
        elif re.search(r"^fait à\s*:", lower):
            print(f"value: {value}, label: {label}, lower: {lower}")
            if value:
                synthesis["location"] = value

        elif in_experiences and text:
            exp_buffer.append(text)
        # else:
            # print(f"value: {value}, label: {label}, lower: {lower}")

    # flush experience buffer
    if exp_buffer:
        background["professional_experiences"].append({"responsibilities": " ".join(exp_buffer)})

    def drop_none(d):
        """Recursively remove keys with None values from a dict."""
        return {k: drop_none(v) if isinstance(v, dict) else v for k, v in d.items() if v is not None}

    identity = drop_none(identity)
    profile = drop_none(profile)

    # Required identity fields must exist
    identity.setdefault("full_name", "")
    identity.setdefault("email", "")
    identity.setdefault("phone", "")

    doc_obj = {
        "_id": str(uuid.uuid4()),
        "candidate_id": str(uuid.uuid4()),
        "status": "SEEKING",
        "tp_type": tp_type,
        "created_at": datetime.utcnow(),
        "identity": identity,
    }

    if any(v for v in education.values() if v):
        doc_obj["education"] = education
    if training_site:
        doc_obj["training_site"] = training_site
    if any(v for v in support.values() if v is not None):
        doc_obj["support"] = support
    if immersion_agreement is not None:
        doc_obj["immersion_agreement"] = immersion_agreement
    if background["last_diploma"] if "last_diploma" in background else background["professional_experiences"]:
        doc_obj["background"] = background
    if any(v for v in profile.values() if v):
        doc_obj["profile"] = profile
    if any(v for v in professional_projects.values() if v):
        doc_obj["professional_projects"] = professional_projects

    skills = extract_skills(doc)
    if skills:
        doc_obj["skills_assessment"] = skills

    sectors, company_skills = extract_sectors_and_company_skills(doc)
    if sectors:
        doc_obj["desired_sectors"] = sectors
    if company_skills:
        doc_obj["expected_company_skills"] = company_skills

    if any(v for v in job_info.values() if v is not None):
        doc_obj["job_info"] = job_info
    if any(v for v in synthesis.values() if v and v != {}):
        doc_obj["synthesis"] = synthesis

    return doc_obj


def insert_docx(file_path, collection):
    candidate = parse_docx(file_path)
    result = collection.insert_one(candidate)
    print(f"Inserted {os.path.basename(file_path)} → tp_type={candidate['tp_type']}  _id={result.inserted_id}")


if __name__ == "__main__":
    collection = get_mongo_collection()

    if len(sys.argv) > 1:
        files = sys.argv[1:]
    else:
        ressources_dir = os.path.abspath(
            os.path.join(os.path.dirname(__file__), "..", "ressources")
        )
        files = sorted(
            os.path.join(ressources_dir, f)
            for f in os.listdir(ressources_dir)
            if f.startswith("ab_apprenti") and f.endswith(".docx")
        )

    inserted, errors = 0, 0
    for path in files:
        if path == '--print':
            continue
        try:
            if "--print" in sys.argv:
                candidate = parse_docx(path)
                print(candidate)
            else:
                insert_docx(path, collection)
                inserted += 1
        except Exception as e:
            print(f"Error processing {path}: {e}")
            errors += 1

    print(f"\nDone: {inserted} inserted, {errors} errors")
